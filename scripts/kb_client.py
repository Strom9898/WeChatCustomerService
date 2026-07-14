#!/usr/bin/env python3
"""
快亮家涂装知识库（零依赖版）
纯 numpy + scipy + jieba 分词，不下载任何模型
"""
import os
import json
import re
import math
from collections import Counter
from typing import List, Dict, Optional

import numpy as np
from scipy.sparse import csr_matrix
from scipy.spatial.distance import cosine

KB_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "knowledge_base")
DOCS = [
    ("用户旅程手册", os.path.expanduser("~/Desktop/快亮家《用户旅程手册》1.0版本_副本.docx")),
    ("文化手册", os.path.expanduser("~/Desktop/快亮家文化手册编写1124_副本.docx")),
]


def extract_text(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- 表格 {i+1} ---")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def chunk_text(text: str, source: str, max_chars: int = 400) -> List[Dict]:
    chunks = []
    current_section = "概述"
    buffer = ""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        title = re.match(r'^第[一二三四五六七八九十]+[章节]', line)
        subtitle = re.match(r'^[0-9]+\.[0-9]', line)
        if title or subtitle:
            if buffer.strip():
                chunks.append({"text": buffer.strip(), "source": source, "section": current_section})
                buffer = ""
            current_section = line[:60]
        buffer += line + "\n"
        if len(buffer) >= max_chars:
            chunks.append({"text": buffer.strip(), "source": source, "section": current_section})
            buffer = ""
    if buffer.strip():
        chunks.append({"text": buffer.strip(), "source": source, "section": current_section})
    return chunks


def tokenize(text: str) -> List[str]:
    """中文分词 + 去停用词"""
    text = text.lower()
    # 按非字母数字拆分（适用于中英文混合）
    tokens = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z0-9]+', text)
    stopwords = set("的了在是和我有不这个人一个上有来大们也到时会就可以对".split())
    return [t for t in tokens if len(t) >= 2 and t not in stopwords]


class KnowledgeBase:
    """TF-IDF 向量知识库（零外部依赖）"""

    def __init__(self):
        self.chunks: List[Dict] = []
        self.vocab: Dict[str, int] = {}
        self.idf: Dict[str, float] = {}
        self.tfidf_matrix: Optional[csr_matrix] = None
        self._ready = False

    def build(self):
        """构建知识库"""
        os.makedirs(KB_DIR, exist_ok=True)
        all_chunks = []

        for src_name, doc_path in DOCS:
            if not os.path.exists(doc_path):
                print(f"⚠️ 文件不存在: {doc_path}")
                continue
            print(f"📖 读取: {src_name}...")
            text = extract_text(doc_path)
            chunks = chunk_text(text, src_name)
            print(f"   → {len(chunks)} 个片段")
            all_chunks.extend(chunks)

        if not all_chunks:
            print("❌ 未读取到内容")
            return False

        print(f"\n📊 构建 TF-IDF 索引（共 {len(all_chunks)} 段）...")
        self.chunks = all_chunks

        # 1. 分词
        tokenized = [tokenize(c["text"]) for c in self.chunks]

        # 2. 构建词汇表
        all_tokens = set()
        for tokens in tokenized:
            all_tokens.update(tokens)
        self.vocab = {t: i for i, t in enumerate(sorted(all_tokens))}
        V = len(self.vocab)
        N = len(self.chunks)

        print(f"   📝 词汇量: {V} 个")

        # 3. 计算 TF-IDF
        doc_freq = Counter()
        for tokens in tokenized:
            doc_freq.update(set(tokens))

        self.idf = {}
        for token, freq in doc_freq.items():
            self.idf[token] = math.log((N + 1) / (freq + 1)) + 1

        # 4. 构建稀疏矩阵
        rows, cols, data = [], [], []
        for i, tokens in enumerate(tokenized):
            tf = Counter(tokens)
            max_tf = max(tf.values()) if tf else 1
            for token, count in tf.items():
                if token in self.vocab:
                    tfidf = (count / max_tf) * self.idf.get(token, 1)
                    rows.append(i)
                    cols.append(self.vocab[token])
                    data.append(tfidf)

        self.tfidf_matrix = csr_matrix((data, (rows, cols)), shape=(N, V))
        self._save()
        print(f"✅ 知识库构建完成！路径: {KB_DIR}")
        return True

    def _encode_query(self, query: str) -> np.ndarray:
        """将查询文本转为 TF-IDF 向量"""
        tokens = tokenize(query)
        if not tokens:
            return np.zeros(len(self.vocab))
        tf = Counter(tokens)
        max_tf = max(tf.values())
        vec = np.zeros(len(self.vocab))
        for token, count in tf.items():
            if token in self.vocab:
                vec[self.vocab[token]] = (count / max_tf) * self.idf.get(token, 1)
        return vec

    def search(self, query: str, n: int = 3) -> List[Dict]:
        """检索最相关知识片段"""
        if not self._ready or self.tfidf_matrix is None:
            return []

        q_vec = self._encode_query(query)
        if q_vec.sum() == 0:
            return []

        # 余弦相似度
        q_norm = q_vec / (np.linalg.norm(q_vec) + 1e-10)
        scores = self.tfidf_matrix.dot(q_norm).toarray().flatten()
        top_k = np.argsort(scores)[-n:][::-1]

        results = []
        for idx in top_k:
            if scores[idx] > 0:
                results.append({
                    "text": self.chunks[idx]["text"][:500],
                    "source": self.chunks[idx]["source"],
                    "section": self.chunks[idx]["section"][:60],
                    "score": round(float(scores[idx]), 4),
                })
        return results

    def format_context(self, query: str, max_chars: int = 1500) -> str:
        """检索并格式化为上下文"""
        chunks = self.search(query, n=3)
        if not chunks:
            return ""
        ctx = "\n\n【相关知识】\n"
        total = 0
        for i, c in enumerate(chunks, 1):
            text = c["text"]
            if total + len(text) > max_chars:
                text = text[:max_chars - total] + "..."
            ctx += f"\n[{i}]({c['source']}-{c['section']}, 匹配度{c['score']}):\n{text}\n"
            total += len(text)
            if total >= max_chars:
                break
        return ctx

    def is_ready(self) -> bool:
        return self._ready

    def _save(self):
        with open(os.path.join(KB_DIR, "chunks.json"), "w", encoding="utf-8") as f:
            json.dump(self.chunks, f, ensure_ascii=False)
        with open(os.path.join(KB_DIR, "vocab.json"), "w", encoding="utf-8") as f:
            json.dump(self.vocab, f, ensure_ascii=False)
        with open(os.path.join(KB_DIR, "idf.json"), "w", encoding="utf-8") as f:
            json.dump(self.idf, f, ensure_ascii=False)
        # 保存 TF-IDF 矩阵
        from scipy.sparse import save_npz
        save_npz(os.path.join(KB_DIR, "tfidf.npz"), self.tfidf_matrix)

    def load(self) -> bool:
        """加载已有知识库"""
        try:
            from scipy.sparse import load_npz
            with open(os.path.join(KB_DIR, "chunks.json"), encoding="utf-8") as file:
                self.chunks = json.load(file)
            with open(os.path.join(KB_DIR, "vocab.json"), encoding="utf-8") as file:
                self.vocab = json.load(file)
            with open(os.path.join(KB_DIR, "idf.json"), encoding="utf-8") as file:
                self.idf = json.load(file)
            self.tfidf_matrix = load_npz(os.path.join(KB_DIR, "tfidf.npz"))
            self._ready = True
            return True
        except Exception as e:
            print(f"Knowledge base load failed: {e}")
            return False


if __name__ == "__main__":
    import sys
    kb = KnowledgeBase()
    if len(sys.argv) > 1 and sys.argv[1] == "build":
        kb.build()
    elif kb.load():
        print(f"✅ 知识库已加载（{len(kb.chunks)} 条）")
        if len(sys.argv) > 1:
            query = " ".join(sys.argv[1:])
            print(f"\n🔍 查询: {query}")
            for r in kb.search(query):
                print(f"\n  [{r['score']}] ({r['source']}-{r['section']})")
                print(f"  {r['text'][:200]}")
    else:
        print("❌ 知识库未构建，请运行: python3 scripts/kb_client.py build")
